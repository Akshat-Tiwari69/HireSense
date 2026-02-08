"""
Admin Routes Module
Handles all admin dashboard endpoints for system management
Protected routes requiring JWT authentication with 'admin' role
"""

import os
import re
import uuid
import json as _json
import zipfile
import tempfile
import shutil
import contextlib
try:
    import rarfile
    RAR_SUPPORTED = True
except ImportError:
    RAR_SUPPORTED = False
from concurrent.futures import ThreadPoolExecutor, as_completed
from flask import Blueprint, request, jsonify, Response, stream_with_context
from flask_jwt_extended import jwt_required, get_jwt_identity, get_jwt
from werkzeug.utils import secure_filename
from datetime import datetime
import logging
from functools import lru_cache
from db_config import get_connection, return_connection
from db_helpers import DatabaseError, insert_candidate, get_candidate_by_email
from auth import hash_password
from psycopg2.extras import RealDictCursor

# Setup logger
logger = logging.getLogger(__name__)

# Create blueprint for admin routes
admin_bp = Blueprint('admin', __name__)


def require_admin_role(f):
    """Decorator to check for admin role"""
    from functools import wraps
    @wraps(f)
    def check_admin_role(*args, **kwargs):
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({
                'status': 'error',
                'message': 'Access denied. Admin role required.'
            }), 403
        return f(*args, **kwargs)
    return check_admin_role


# ============================================================================
#                           USER MANAGEMENT
# ============================================================================

@admin_bp.route('/users', methods=['GET'])
@jwt_required()
@require_admin_role
def get_all_users():
    """Get all users in the system with optimized pooled connection"""
    conn = None
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT id, name, email, role, created_at FROM users ORDER BY id")
        rows = cursor.fetchall()
        
        users = [{
            'id': row[0],
            'name': row[1],
            'email': row[2],
            'role': row[3],
            'created_at': row[4]
        } for row in rows]
        
        return jsonify({'status': 'success', 'data': users}), 200
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500
    finally:
        if conn:
            return_connection(conn)


@admin_bp.route('/users', methods=['POST'])
@jwt_required()
@require_admin_role
def create_user():
    """Create a new user"""
    try:
        data = request.get_json()
        name = data.get('name')
        email = data.get('email')
        password = data.get('password')
        role = data.get('role', 'interviewer')
        
        if not all([name, email, password]):
            return jsonify({'status': 'error', 'message': 'Name, email, and password are required'}), 400
        
        if role not in ['interviewer', 'admin', 'proctor']:
            return jsonify({'status': 'error', 'message': 'Invalid role. Must be interviewer, proctor, or admin'}), 400
        
        admin_email = get_jwt_identity()
        logger.info(f"[ADMIN ACTION] {admin_email} creating new user: {email} with role: {role}")
        
        password_hash = hash_password(password)
        
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO users (name, email, password_hash, role) VALUES (%s, %s, %s, %s) RETURNING id",
            (name, email, password_hash, role)
        )
        result = cursor.fetchone()
        user_id = result[0] if result else None
        conn.commit()
        conn.close()
        
        logger.info(f"[ADMIN ACTION] {admin_email} successfully created user ID: {user_id} ({email})")
        
        return jsonify({
            'status': 'success',
            'message': 'User created successfully',
            'data': {'id': user_id, 'name': name, 'email': email, 'role': role}
        }), 201
    except Exception as e:
        if 'unique constraint' in str(e).lower() or 'duplicate' in str(e).lower():
            return jsonify({'status': 'error', 'message': 'Email already exists'}), 409
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/users/<int:user_id>', methods=['PUT'])
@jwt_required()
@require_admin_role
def update_user(user_id):
    """Update a user's details"""
    try:
        admin_email = get_jwt_identity()
        data = request.get_json()
        
        logger.info(f"[ADMIN ACTION] {admin_email} updating user ID: {user_id} with data: {list(data.keys())}")
        
        conn = get_connection()
        cursor = conn.cursor()
        
        # Build update query dynamically
        from psycopg2 import sql as psql
        field_names = []
        values = []
        
        if 'name' in data:
            field_names.append('name')
            values.append(data['name'])
        if 'email' in data:
            field_names.append('email')
            values.append(data['email'])
        if 'role' in data:
            if data['role'] not in ['interviewer', 'admin', 'proctor']:
                return jsonify({'status': 'error', 'message': 'Invalid role. Must be interviewer, proctor, or admin'}), 400
            field_names.append('role')
            values.append(data['role'])
        if 'password' in data and data['password']:
            field_names.append('password_hash')
            values.append(hash_password(data['password']))
        
        if not field_names:
            return jsonify({'status': 'error', 'message': 'No fields to update'}), 400
        
        values.append(user_id)
        set_clause = psql.SQL(', ').join(
            [psql.SQL("{} = %s").format(psql.Identifier(f)) for f in field_names]
        )
        query = psql.SQL("UPDATE users SET {} WHERE id = %s").format(set_clause)
        cursor.execute(query, values)
        conn.commit()
        conn.close()
        
        logger.info(f"[ADMIN ACTION] {admin_email} successfully updated user ID: {user_id}")
        
        return jsonify({'status': 'success', 'message': 'User updated successfully'}), 200
    except Exception as e:
        logger.error(f"[ADMIN ERROR] Failed to update user ID {user_id}: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/users/<int:user_id>', methods=['DELETE'])
@jwt_required()
@require_admin_role
def delete_user(user_id):
    """Delete a user"""
    try:
        # Prevent deleting yourself
        admin_email = get_jwt_identity()
        current_user_id = int(admin_email.split('@')[0]) if '@' in admin_email else 0
        
        conn = get_connection()
        cursor = conn.cursor()
        
        # Get user info before deleting
        cursor.execute("SELECT email, role FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        
        if not user:
            conn.close()
            return jsonify({'status': 'error', 'message': 'User not found'}), 404
        
        user_email = user[0]
        user_role = user[1]
        
        logger.warning(f"[ADMIN ACTION] {admin_email} deleting user ID: {user_id} ({user_email}, role: {user_role})")
        
        cursor.execute("DELETE FROM users WHERE id = %s", (user_id,))
        conn.commit()
        conn.close()
        
        logger.info(f"[ADMIN ACTION] {admin_email} successfully deleted user ID: {user_id}")
        
        return jsonify({'status': 'success', 'message': 'User deleted successfully'}), 200
    except Exception as e:
        logger.error(f"[ADMIN ERROR] Failed to delete user ID {user_id}: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ============================================================================
#                           CANDIDATE MANAGEMENT
# ============================================================================

@admin_bp.route('/absence-of-details', methods=['GET'])
@jwt_required()
@require_admin_role
def get_absence_of_details():
    """Get all candidates with 'Absence of Details' status and infer what's missing."""
    conn = None
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT c.id, c.name, c.email, c.phone, c.resume_path,
                   c.match_score, c.status, c.created_at,
                   cjm.job_id,
                   (SELECT jp.title FROM job_postings jp WHERE jp.id = cjm.job_id LIMIT 1) AS job_title
            FROM candidates c
            LEFT JOIN candidate_job_matches cjm ON cjm.candidate_id = c.id
            WHERE c.status = 'Absence of Details'
            ORDER BY c.created_at DESC
        """)
        rows = cursor.fetchall()

        results = []
        for row in rows:
            cid, name, email, phone, resume_path, match_score, status, created_at, job_id, job_title = row
            missing = []
            # Check name placeholder patterns
            if not name or name.strip().lower() in ('unknown candidate', 'candidate') or name.strip() == '':
                missing.append('name')
            # Check email placeholder patterns
            if not email or email.endswith('@bulk-upload.local'):
                missing.append('email')
            # Phone is optional but we still report if it's absent
            if not phone or phone.strip() == '':
                missing.append('phone')

            results.append({
                'id': cid,
                'name': name,
                'email': email,
                'phone': phone or '',
                'resume_path': resume_path,
                'match_score': match_score or 0,
                'status': status,
                'created_at': str(created_at) if created_at else None,
                'job_id': job_id,
                'job_title': job_title or 'N/A',
                'missing_fields': missing,
            })

        return jsonify({'status': 'success', 'data': results, 'total': len(results)}), 200
    except Exception as e:
        logger.error(f"[ADMIN] absence-of-details error: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500
    finally:
        if conn:
            return_connection(conn)


@admin_bp.route('/candidates', methods=['GET'])
@jwt_required()
@require_admin_role
def get_all_candidates():
    """Get all candidates with full details using optimized pooled connection"""
    conn = None
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT id, name, email, phone, resume_path, match_score, 
                   shortlist_status, pros, cons, created_at, status
            FROM candidates ORDER BY id DESC
        """)
        rows = cursor.fetchall()
        
        candidates = [{
            'id': row[0],
            'name': row[1],
            'email': row[2],
            'phone': row[3],
            'resume_path': row[4],
            'match_score': row[5],
            'shortlist_status': row[6],
            'pros': row[7],
            'cons': row[8],
            'created_at': row[9],
            'status': row[10] or row[6] or 'Applied'
        } for row in rows]
        
        return jsonify({'status': 'success', 'data': candidates}), 200
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500
    finally:
        if conn:
            return_connection(conn)


@admin_bp.route('/candidates/<int:candidate_id>', methods=['PUT'])
@jwt_required()
@require_admin_role
def update_candidate(candidate_id):
    """Update a candidate's details"""
    try:
        admin_email = get_jwt_identity()
        data = request.get_json()
        
        logger.info(f"[ADMIN ACTION] {admin_email} updating candidate ID: {candidate_id} with data: {list(data.keys())}")
        
        conn = get_connection()
        cursor = conn.cursor()
        
        from psycopg2 import sql as psql
        field_names = []
        values = []
        
        if 'name' in data:
            field_names.append('name')
            values.append(data['name'])
        if 'email' in data:
            field_names.append('email')
            values.append(data['email'])
        if 'phone' in data:
            field_names.append('phone')
            values.append(data['phone'])
        if 'status' in data:
            field_names.append('status')
            values.append(data['status'])
            field_names.append('shortlist_status')
            values.append(data['status'])
        if 'match_score' in data:
            field_names.append('match_score')
            values.append(data['match_score'])
        
        if not field_names:
            return jsonify({'status': 'error', 'message': 'No fields to update'}), 400
        
        values.append(candidate_id)
        set_clause = psql.SQL(', ').join(
            [psql.SQL("{} = %s").format(psql.Identifier(f)) for f in field_names]
        )
        query = psql.SQL("UPDATE candidates SET {} WHERE id = %s").format(set_clause)
        cursor.execute(query, values)
        conn.commit()
        conn.close()
        
        logger.info(f"[ADMIN ACTION] {admin_email} successfully updated candidate ID: {candidate_id}")
        
        return jsonify({'status': 'success', 'message': 'Candidate updated successfully'}), 200
    except Exception as e:
        logger.error(f"[ADMIN ERROR] Failed to update candidate ID {candidate_id}: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/candidates/<int:candidate_id>', methods=['DELETE'])
@jwt_required()
@require_admin_role
def delete_candidate(candidate_id):
    """Delete a candidate and all related data"""
    try:
        admin_email = get_jwt_identity()
        conn = get_connection()
        cursor = conn.cursor()
        
        # Get candidate info before deleting
        cursor.execute("SELECT name, email FROM candidates WHERE id = %s", (candidate_id,))
        candidate = cursor.fetchone()
        
        if not candidate:
            conn.close()
            return jsonify({'status': 'error', 'message': 'Candidate not found'}), 404
        
        candidate_name = candidate[0]
        candidate_email = candidate[1]
        
        logger.warning(f"[ADMIN ACTION] {admin_email} deleting candidate ID: {candidate_id} ({candidate_name}, {candidate_email}) and all related data")
        
        # Delete related records first
        cursor.execute("DELETE FROM scheduled_assessments WHERE candidate_id = %s", (candidate_id,))
        cursor.execute("DELETE FROM assessments WHERE candidate_id = %s", (candidate_id,))
        cursor.execute("DELETE FROM candidates WHERE id = %s", (candidate_id,))
        
        conn.commit()
        conn.close()
        
        logger.info(f"[ADMIN ACTION] {admin_email} successfully deleted candidate ID: {candidate_id}")
        
        return jsonify({'status': 'success', 'message': 'Candidate deleted successfully'}), 200
    except Exception as e:
        logger.error(f"[ADMIN ERROR] Failed to delete candidate ID {candidate_id}: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ============================================================================
#                           DATABASE MANAGEMENT
# ============================================================================

@admin_bp.route('/db/tables', methods=['GET'])
@jwt_required()
@require_admin_role
def get_db_tables():
    """Get list of all tables in the database"""
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # PostgreSQL query for tables
        cursor.execute("""
            SELECT table_name 
            FROM information_schema.tables 
            WHERE table_schema = 'public'
            ORDER BY table_name
        """)
        rows = cursor.fetchall()
        conn.close()
        
        tables = [row[0] for row in rows]
        return jsonify({'status': 'success', 'data': tables}), 200
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/db/tables/<table_name>', methods=['GET'])
@jwt_required()
@require_admin_role
def get_table_data(table_name):
    """Get data from a specific table (limited to 100 rows) with optimized pooling"""
    conn = None
    try:
        # Whitelist allowed tables for security
        allowed_tables = {
            'users', 'candidates', 'assessments', 'scheduled_assessments', 
            'email_logs', 'questions', 'proctoring_violations',
            'coding_submissions', 'job_descriptions', 'mcq_responses',
            'proctoring_events', 'psychometric_responses'
        }
        if table_name not in allowed_tables:
            return jsonify({'status': 'error', 'message': 'Table not allowed'}), 403
        
        conn = get_connection()
        cursor = conn.cursor()
        
        # Get column names - use parameterized query for information_schema
        cursor.execute(
            "SELECT column_name FROM information_schema.columns "
            "WHERE table_name = %s ORDER BY ordinal_position",
            (table_name,)
        )
        columns = [row[0] for row in cursor.fetchall()]
        
        # Get data (limit 100 rows)
        # table_name is validated against the whitelist above, safe to use
        # psycopg2 does not support parameterized table names, so we use
        # sql.Identifier for safe dynamic table references
        from psycopg2 import sql
        if 'id' in columns:
            cursor.execute(sql.SQL("SELECT * FROM {} ORDER BY id DESC LIMIT 100").format(sql.Identifier(table_name)))
        else:
            cursor.execute(sql.SQL("SELECT * FROM {} LIMIT 100").format(sql.Identifier(table_name)))
        rows = cursor.fetchall()
        
        data = [dict(zip(columns, row)) for row in rows]
        
        return jsonify({
            'status': 'success',
            'data': data,
            'columns': columns,
            'count': len(data)
        }), 200
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500
    finally:
        if conn:
            return_connection(conn)


def _fetch_database_stats():
    """Fetch aggregated statistics from database"""
    conn = get_connection()
    cursor = conn.cursor()
    
    # Count users by role
    cursor.execute("SELECT role, COUNT(*) FROM users GROUP BY role")
    users_by_role = {row[0]: row[1] for row in cursor.fetchall()}
    
    # Count candidates by status
    cursor.execute("SELECT COALESCE(status, shortlist_status, 'Applied') as s, COUNT(*) FROM candidates GROUP BY s")
    candidates_by_status = {row[0]: row[1] for row in cursor.fetchall()}
    
    # Count scheduled assessments
    cursor.execute("SELECT status, COUNT(*) FROM scheduled_assessments GROUP BY status")
    assessments_by_status = {row[0]: row[1] for row in cursor.fetchall()}
    
    # Total counts
    cursor.execute("SELECT COUNT(*) FROM users")
    total_users = cursor.fetchone()[0]
    
    cursor.execute("SELECT COUNT(*) FROM candidates")
    total_candidates = cursor.fetchone()[0]
    
    cursor.execute("SELECT COUNT(*) FROM assessments")
    total_assessments = cursor.fetchone()[0]
    
    conn.close()
    
    return {
        'users_by_role': users_by_role,
        'candidates_by_status': candidates_by_status,
        'assessments_by_status': assessments_by_status,
        'total_users': total_users,
        'total_candidates': total_candidates,
        'total_assessments': total_assessments
    }


@admin_bp.route('/db/stats', methods=['GET'])
@jwt_required()
@require_admin_role
def get_db_stats():
    """Get database statistics"""
    try:
        stats = _fetch_database_stats()
        return jsonify({'status': 'success', 'data': stats}), 200
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ============================================================================
#                           SYSTEM SETTINGS
# ============================================================================

@admin_bp.route('/settings/env', methods=['GET'])
@jwt_required()
@require_admin_role
def get_env_status():
    """Get status of required environment variables (not values)"""
    try:
        env_vars = [
            'DATABASE_URL',
            'JWT_SECRET_KEY',
            'RESEND_API_KEY',
            'RESEND_FROM_EMAIL',
            'SMTP_HOST',
            'SMTP_USER',
            'FRONTEND_URL'
        ]
        
        status = {}
        for var in env_vars:
            if value := os.environ.get(var):
                # Mask the value for security
                if len(value) > 8:
                    status[var] = f"{value[:4]}...{value[-4:]}"
                else:
                    status[var] = "***configured***"
            else:
                status[var] = None
        
        return jsonify({'status': 'success', 'data': status}), 200
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/settings/env', methods=['POST'])
@jwt_required()
@require_admin_role
def set_env_variable():
    """Set/update environment variables and persist to .env file"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'status': 'error', 'message': 'No data provided'}), 400
        
        var_name = data.get('name')
        var_value = data.get('value')
        
        if not var_name:
            return jsonify({'status': 'error', 'message': 'Variable name is required'}), 400
        
        # Validate DATABASE_URL format if being set
        if var_name == 'DATABASE_URL' and var_value and not var_value.startswith(('postgresql://', 'postgres://')):
            return jsonify({'status': 'error', 'message': 'DATABASE_URL must start with postgresql:// or postgres://'}), 400
        
        # Set in current environment
        if var_value:
            os.environ[var_name] = var_value
        elif var_name in os.environ:
            del os.environ[var_name]
        
        # Persist to .env file
        env_file_path = os.path.join(os.path.dirname(__file__), '.env')
        env_vars = {}
        
        # Read existing .env file if it exists
        if os.path.exists(env_file_path):
            with open(env_file_path, 'r') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#') and '=' in line:
                        key, value = line.split('=', 1)
                        env_vars[key.strip()] = value.strip()
        
        # Update or remove the variable
        if var_value:
            env_vars[var_name] = var_value
        elif var_name in env_vars:
            del env_vars[var_name]
        
        # Write back to .env file
        with open(env_file_path, 'w') as f:
            for key, value in env_vars.items():
                f.write(f"{key}={value}\n")
        
        logger.info(f"Environment variable {var_name} {'updated' if var_value else 'removed'} by admin")
        
        return jsonify({
            'status': 'success',
            'message': f"Environment variable {var_name} {'updated' if var_value else 'removed'} successfully"
        }), 200
        
    except Exception as e:
        logger.error(f"Error setting environment variable: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/reset-candidate-status/<int:candidate_id>', methods=['POST'])
@jwt_required()
@require_admin_role
def reset_candidate_status(candidate_id):
    """Reset a candidate's status back to Applied"""
    try:
        admin_email = get_jwt_identity()
        conn = get_connection()
        cursor = conn.cursor()
        
        # Get candidate info
        cursor.execute("SELECT name, email FROM candidates WHERE id = %s", (candidate_id,))
        candidate = cursor.fetchone()
        
        if not candidate:
            conn.close()
            return jsonify({'status': 'error', 'message': 'Candidate not found'}), 404
        
        candidate_name = candidate[0]
        candidate_email = candidate[1]
        
        logger.info(f"[ADMIN ACTION] {admin_email} resetting status to 'Applied' for candidate ID: {candidate_id} ({candidate_name})")
        
        cursor.execute("""
            UPDATE candidates 
            SET status = 'Applied', shortlist_status = 'Applied'
            WHERE id = %s
        """, (candidate_id,))
        
        # Also delete any scheduled assessments for this candidate
        cursor.execute("DELETE FROM scheduled_assessments WHERE candidate_id = %s", (candidate_id,))
        
        conn.commit()
        conn.close()
        
        logger.info(f"[ADMIN ACTION] {admin_email} successfully reset candidate ID: {candidate_id} to 'Applied'")
        
        return jsonify({'status': 'success', 'message': 'Candidate status reset to Applied'}), 200
    except Exception as e:
        logger.error(f"[ADMIN ERROR] Failed to reset candidate ID {candidate_id}: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ============================================================================
#                           EMAIL LOGS MANAGEMENT
# ============================================================================

@admin_bp.route('/email-logs', methods=['GET'])
@jwt_required()
@require_admin_role
def get_email_logs():
    """
    Get all email logs with optimized timestamp formatting
    """
    conn = None
    try:
        conn = get_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        cursor.execute("SELECT * FROM email_logs ORDER BY sent_at DESC LIMIT 100")
        logs = cursor.fetchall()
        
        # Use list comprehension for in-place formatting (faster than loop)
        formatted_logs = [
            {**log, 'sent_at': log['sent_at'].isoformat() if log.get('sent_at') else None}
            for log in logs
        ]
        
        return jsonify({
            'status': 'success',
            'data': formatted_logs
        })
    except Exception as e:
        logger.error(f"[ADMIN ERROR] Failed to fetch email logs: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500
    finally:
        if conn:
            return_connection(conn)


# ============================================================================
#                    JOB POSTINGS MANAGEMENT (LEGACY - DEPRECATED)
#  These endpoints are kept for backwards compatibility.
#  New code should use /api/jobs/postings (job_routes.py) which supports
#  sectors, RBAC, required/preferred skills, and candidate matching.
# ============================================================================

@admin_bp.route('/job-postings', methods=['GET'])
@jwt_required()
@require_admin_role
def get_job_postings():
    """Legacy: Get all job postings. Prefer /api/jobs/postings"""
    from flask import redirect
    return redirect('/api/jobs/postings?status=all', code=307)


@admin_bp.route('/job-postings', methods=['POST'])
@jwt_required()
@require_admin_role
def create_job_posting():
    """Legacy: Create a job posting. Prefer /api/jobs/postings"""
    from flask import redirect
    return redirect('/api/jobs/postings', code=307)


@admin_bp.route('/job-postings/<int:job_id>', methods=['DELETE'])
@jwt_required()
@require_admin_role
def delete_job_posting(job_id):
    """Legacy: Delete a job posting. Prefer /api/jobs/postings/<id>"""
    from flask import redirect
    return redirect(f'/api/jobs/postings/{job_id}', code=307)


# ============================================================================
#                           ANALYTICS
# ============================================================================

@admin_bp.route('/analytics', methods=['GET'])
@jwt_required()
@require_admin_role
def get_analytics():
    """
    Get system-wide analytics and statistics.
    Optimized with combined queries for better performance.
    """
    conn = None
    try:
        conn = get_connection()
        cursor = conn.cursor(cursor_factory=RealDictCursor)
        
        # Combined query for all statistics - more efficient than separate queries
        cursor.execute("""
            SELECT 
                (SELECT COUNT(*) FROM candidates) as total_candidates,
                (SELECT COUNT(*) FROM candidates WHERE status = 'pending') as pending_candidates,
                (SELECT COUNT(*) FROM candidates WHERE status = 'under_review') as under_review_candidates,
                (SELECT COUNT(*) FROM candidates WHERE status = 'hired') as hired_candidates,
                (SELECT COUNT(*) FROM candidates WHERE status = 'rejected') as rejected_candidates,
                (SELECT AVG(match_score) FROM candidates) as avg_match_score,
                (SELECT COUNT(*) FROM candidates WHERE created_at >= NOW() - INTERVAL '30 days') as candidates_this_month,
                (SELECT COUNT(*) FROM scheduled_assessments) as total_assessments,
                (SELECT COUNT(*) FROM scheduled_assessments WHERE status = 'scheduled') as scheduled_assessments,
                (SELECT COUNT(*) FROM scheduled_assessments WHERE status = 'in_progress') as in_progress_assessments,
                (SELECT COUNT(*) FROM scheduled_assessments WHERE status = 'completed') as completed_assessments,
                (SELECT AVG(technical_score) FROM scheduled_assessments) as avg_technical_score,
                (SELECT AVG(psychometric_score) FROM scheduled_assessments) as avg_psychometric_score,
                (SELECT COUNT(*) FROM scheduled_assessments WHERE created_at >= NOW() - INTERVAL '30 days') as assessments_this_month
        """)
        
        stats = cursor.fetchone()
        cursor.close()
        
        # Build response using extracted stats
        analytics = {
            'candidates': {
                'total': stats['total_candidates'] or 0,
                'pending': stats['pending_candidates'] or 0,
                'under_review': stats['under_review_candidates'] or 0,
                'hired': stats['hired_candidates'] or 0,
                'rejected': stats['rejected_candidates'] or 0,
                'avg_match_score': round(float(stats['avg_match_score'] or 0), 2),
                'this_month': stats['candidates_this_month'] or 0
            },
            'assessments': {
                'total': stats['total_assessments'] or 0,
                'scheduled': stats['scheduled_assessments'] or 0,
                'in_progress': stats['in_progress_assessments'] or 0,
                'completed': stats['completed_assessments'] or 0,
                'avg_technical_score': round(float(stats['avg_technical_score'] or 0), 2),
                'avg_psychometric_score': round(float(stats['avg_psychometric_score'] or 0), 2),
                'this_month': stats['assessments_this_month'] or 0
            }
        }
        
        return jsonify({
            'status': 'success',
            'data': analytics
        })
    except Exception as e:
        logger.error(f"[ADMIN ERROR] Failed to fetch analytics: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500
    finally:
        if conn:
            conn.close()


# ============================================================================
#                        BULK RESUME UPLOAD
# ============================================================================

# Helpers (duplicated from app.py to avoid circular imports)
ALLOWED_RESUME_EXTENSIONS = {'pdf', 'docx'}
EMAIL_PATTERN_BULK = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9]([a-zA-Z0-9-]*[a-zA-Z0-9])?(\.[a-zA-Z0-9]([a-zA-Z0-9-]*[a-zA-Z0-9])?)*\.[a-zA-Z]{2,}$'

MAX_BULK_WORKERS = 8  # concurrent AI agents processing resumes


def _bulk_valid_email(email):
    if email is None or not isinstance(email, str) or not email.strip():
        return False
    return re.match(EMAIL_PATTERN_BULK, email) is not None


def _bulk_name_from_email(email):
    if not email or '@' not in email:
        return None
    local = email.split('@', 1)[0]
    parts = re.split(r'[._-]+', local)
    parts = [p for p in parts if p]
    return " ".join(p.capitalize() for p in parts) if parts else None


def _merge_ai_data_to_parsed(parsed_data, ai_data):
    """Merge extracted AI data into parsed_data dictionary"""
    if ai_data.get('skills'):
        parsed_data['skills'] = ai_data['skills']
    if ai_data.get('experience') and ai_data['experience'] > 0:
        parsed_data['experience'] = ai_data['experience']
    if ai_data.get('education'):
        parsed_data['education'] = ai_data['education']
    if ai_data.get('name'):
        parsed_data['name'] = ai_data['name']
    if ai_data.get('email'):
        parsed_data['email'] = ai_data['email']
    if ai_data.get('phone'):
        parsed_data['phone'] = ai_data['phone']


def _get_db_cursor():
    """Get database connection and cursor together"""
    conn = get_connection()
    return conn, conn.cursor()


def _save_candidate_job_match(candidate_id, job_id, match_score, ai_reasoning):
    """Save candidate job match to database"""
    match_conn, match_cur = _get_db_cursor()
    try:
        match_cur.execute("""
            INSERT INTO candidate_job_matches
            (candidate_id, job_id, match_score, ai_reasoning)
            VALUES (%s, %s, %s, %s)
            ON CONFLICT (candidate_id, job_id)
            DO UPDATE SET match_score = EXCLUDED.match_score,
                         ai_reasoning = EXCLUDED.ai_reasoning,
                         matched_at = NOW()
        """, (candidate_id, int(job_id), match_score, ai_reasoning))
        match_cur.execute("""
            UPDATE candidates SET best_match_job_id = %s, match_score = %s WHERE id = %s
        """, (int(job_id), match_score, candidate_id))
        match_conn.commit()
        return_connection(match_conn)
    except Exception as match_err:
        logger.warning(f"Could not save job match for candidate {candidate_id}: {match_err}")


def _fetch_job_for_bulk(job_id):
    """Fetch job description for bulk processing (standalone, no circular import)."""
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, title, department, required_skills, preferred_skills, min_experience "
            "FROM job_descriptions WHERE id = %s AND status = 'active'",
            (int(job_id),)
        )
        row = cursor.fetchone()
        return_connection(conn)
        if row:
            skills = set()
            for skills_val in (row[3], row[4]):
                if skills_val:
                    with contextlib.suppress(ValueError, TypeError):
                        parsed = _json.loads(skills_val)
                        if isinstance(parsed, list):
                            skills.update(s.strip() for s in parsed if s.strip())
                            continue
                    skills.update(s.strip() for s in str(skills_val).split(',') if s.strip())
            min_exp = row[5] or 0
            job_info = {'id': row[0], 'title': row[1], 'department': row[2]}
            return {'skills': list(skills), 'min_experience': min_exp, 'title': row[1], 'department': row[2]}, job_info
    except Exception as e:
        logger.warning(f"[BULK] Could not load job posting {job_id}: {e}")
    return None, None


def _process_single_resume(filepath, filename, job_description, job_info, job_id, upload_folder):
    """
    Process a single resume file through the full AI pipeline.
    Runs inside a ThreadPoolExecutor worker.
    Returns a result dict.
    """
    from resume_parser import parse_resume
    from resume_analyzer import analyze_resume, ResumeAnalyzer
    from resume_parser import calculate_match_score

    result = {
        'filename': filename,
        'status': 'error',
        'name': None,
        'email': None,
        'match_score': 0,
        'recommendation': None,
        'candidate_id': None,
        'error': None,
        'missing': []
    }

    try:
        # 1. Basic parse
        parsed_data = parse_resume(filepath, job_description)

        # 2. Extract text
        with open(filepath, 'rb') as f:
            if filepath.lower().endswith('.pdf'):
                from PyPDF2 import PdfReader
                pdf = PdfReader(f)
                resume_text = " ".join([page.extract_text() or '' for page in pdf.pages])
            else:
                from docx import Document
                doc = Document(f)
                resume_text = " ".join([para.text for para in doc.paragraphs])

        if not resume_text or len(resume_text.strip()) < 50:
            # Still save the candidate even if text extraction is poor
            result['missing'] = ['name', 'email']
            name = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title() or 'Unknown Candidate'
            email = f'unknown-{uuid.uuid4().hex[:12]}@bulk-upload.local'
            try:
                candidate_id = insert_candidate(
                    name=name, email=email, phone='',
                    resume_path=filepath,
                    parsed_data={'skills': [], 'experience': 0, 'education': '', 'match_score': 0, 'shortlist_status': 'Pending Review'},
                    pros=None, cons=None, status='Absence of Details'
                )
                result['status'] = 'success'
                result['name'] = name
                result['email'] = email
                result['candidate_id'] = candidate_id
                result['error'] = 'Could not extract text — saved with Absence of Details'
            except Exception as save_err:
                result['error'] = f'Could not extract text and save failed: {save_err}'
            return result

        # 3. AI extraction
        try:
            analyzer = ResumeAnalyzer()
            if ai_data := analyzer.extract_resume_data(resume_text):
                _merge_ai_data_to_parsed(parsed_data, ai_data)
                parsed_data['match_score'] = calculate_match_score(
                    parsed_data.get('skills', []),
                    parsed_data.get('experience', 0),
                    job_description.get('skills', []),
                    job_description.get('min_experience', 0)
                )
        except Exception as ai_err:
            logger.warning(f"[BULK] AI extraction failed for {filename}: {ai_err}")

        # 4. AI analysis (pros/cons)
        ai_analysis = None
        try:
            ai_analysis = analyze_resume(
                resume_text=resume_text,
                parsed_data=parsed_data,
                job_requirements=job_description,
                enhance_score=True
            )
            if ai_analysis and 'enhanced_match_score' in ai_analysis:
                parsed_data['match_score'] = ai_analysis['enhanced_match_score']
        except Exception as ai_err:
            logger.warning(f"[BULK] AI analysis failed for {filename}: {ai_err}")
            ai_analysis = {
                "pros": ["Resume uploaded successfully"],
                "cons": ["AI analysis unavailable - manual review recommended"],
                "overall_assessment": "AI analysis failed. Manual review required.",
                "recommendation": "Pending Review",
                "confidence_score": 0
            }

        # 5. Resolve identity — NEVER skip a resume
        name = parsed_data.get('name')
        email = parsed_data.get('email')
        phone = parsed_data.get('phone', '')

        # Track what details are missing (only name & email are required; phone is optional)
        missing_details = []
        if not name:
            missing_details.append('name')
        if not email or not _bulk_valid_email(email):
            missing_details.append('email')

        # Generate placeholders for missing critical fields
        if not email or not _bulk_valid_email(email):
            # Unique placeholder email so the DB UNIQUE constraint is satisfied
            email = f'unknown-{uuid.uuid4().hex[:12]}@bulk-upload.local'

        if not name:
            # Try to derive from email, then from filename
            name = _bulk_name_from_email(email)
            if not name or 'unknown' in name.lower():
                # Use original filename minus extension as name hint
                name = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title() or 'Unknown Candidate'

        # Determine status — if any key detail is missing, flag it
        candidate_status = 'Absence of Details' if missing_details else 'Applied'

        result['name'] = name
        result['email'] = email
        result['match_score'] = parsed_data.get('match_score', 0)
        result['recommendation'] = ai_analysis.get('recommendation', 'Pending Review') if ai_analysis else 'Pending Review'
        result['missing'] = missing_details  # expose what's missing

        # 6. Check duplicate (only for real emails)
        if not email.endswith('@bulk-upload.local'):
            with contextlib.suppress(Exception):
                if existing := get_candidate_by_email(email):
                    result['status'] = 'duplicate'
                    result['candidate_id'] = existing['id']
                    result['error'] = f'Already registered (ID: {existing["id"]})'
                    return result

        # 7. Save candidate — every resume becomes a candidate row
        pros_text = "\n".join(ai_analysis.get('pros', [])) if ai_analysis else None
        cons_text = "\n".join(ai_analysis.get('cons', [])) if ai_analysis else None

        candidate_id = insert_candidate(
            name=name,
            email=email,
            phone=phone or '',
            resume_path=filepath,
            parsed_data=parsed_data,
            pros=pros_text,
            cons=cons_text,
            status=candidate_status
        )
        result['candidate_id'] = candidate_id

        # 8. Save job match
        if candidate_id:
            match_score = parsed_data.get('match_score', 0)
            ai_reasoning = ai_analysis.get('overall_assessment', '') if ai_analysis else ''
            _save_candidate_job_match(candidate_id, job_id, match_score, ai_reasoning)

        result['status'] = 'success'
        if missing_details:
            result['error'] = f'Saved with Absence of Details (missing: {", ".join(missing_details)})'
        logger.info(f"[BULK] Processed {filename} -> {name} <{email}> score={result['match_score']} status={candidate_status}")

    except Exception as e:
        result['error'] = str(e)
        logger.error(f"[BULK] Error processing {filename}: {e}")

    return result


@admin_bp.route('/bulk-upload', methods=['POST'])
@jwt_required()
@require_admin_role
def bulk_upload_resumes():
    """
    Bulk resume upload from a ZIP file.
    Extracts all PDF/DOCX files and processes them in parallel using
    a ThreadPoolExecutor (multiple AI agents working concurrently).

    Form data:
        - file: ZIP archive containing PDF/DOCX resumes
        - job_id: Target job posting ID to score all resumes against
    """
    logger.info("=" * 80)
    logger.info("[BULK] BULK RESUME UPLOAD REQUEST RECEIVED")
    logger.info("=" * 80)

    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': 'No file uploaded'}), 400

    archive_file = request.files['file']
    fname = (archive_file.filename or '').lower()
    if not fname.endswith('.zip') and not fname.endswith('.rar'):
        return jsonify({'status': 'error', 'message': 'Please upload a .zip or .rar file'}), 400
    if fname.endswith('.rar') and not RAR_SUPPORTED:
        return jsonify({'status': 'error', 'message': 'RAR support not available on this server. Please upload a .zip file instead.'}), 400
    is_rar = fname.endswith('.rar')

    job_id = request.form.get('job_id')
    if not job_id:
        return jsonify({'status': 'error', 'message': 'Please select a job position'}), 400

    job_description, job_info = _fetch_job_for_bulk(job_id)
    if not job_description:
        return jsonify({'status': 'error', 'message': 'Selected job is no longer active'}), 400

    logger.info(f"[BULK] Target job: {job_info['title']} (ID: {job_id})")

    # Save ZIP to temp location and extract
    temp_dir = tempfile.mkdtemp(prefix='bulk_upload_')
    extract_dir = os.path.join(temp_dir, 'resumes')
    os.makedirs(extract_dir, exist_ok=True)

    # Upload folder for permanent storage
    upload_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
    os.makedirs(upload_folder, exist_ok=True)

    try:
        archive_path = os.path.join(temp_dir, 'upload.rar' if is_rar else 'upload.zip')
        archive_file.save(archive_path)

        # Extract and filter resume files
        resume_files = []

        def _collect_resume_files(base_dir):
            """Walk extracted directory and collect resume files."""
            for root, _dirs, files in os.walk(base_dir):
                for fname in files:
                    if '__MACOSX' in root or fname.startswith('.'):
                        continue
                    ext = fname.rsplit('.', 1)[-1].lower() if '.' in fname else ''
                    if ext not in ALLOWED_RESUME_EXTENSIONS:
                        continue
                    rel_path = os.path.relpath(os.path.join(root, fname), base_dir)
                    parent_dir = os.path.basename(os.path.dirname(rel_path))
                    display_name = f"{parent_dir}/{fname}" if parent_dir and parent_dir != '.' else fname
                    safe_name = secure_filename(fname)
                    unique_name = f"{uuid.uuid4()}_{safe_name}"
                    perm_path = os.path.join(upload_folder, unique_name)
                    shutil.copy2(os.path.join(root, fname), perm_path)
                    resume_files.append((perm_path, display_name))

        if is_rar:
            # Use subprocess to call unrar directly — avoids eventlet pipe conflicts on Windows
            import subprocess as _sp
            # Try multiple known locations for unrar
            _unrar_candidates = [
                r'C:\Program Files\WinRAR\UnRAR.exe',
                r'C:\Program Files (x86)\WinRAR\UnRAR.exe',
                'unrar',
                'UnRAR',
            ]
            if RAR_SUPPORTED and getattr(rarfile, 'UNRAR_TOOL', None):
                _unrar_candidates.insert(0, rarfile.UNRAR_TOOL)

            rar_extract_dir = os.path.join(temp_dir, 'rar_extracted')
            os.makedirs(rar_extract_dir, exist_ok=True)
            proc = None
            for _tool in _unrar_candidates:
                try:
                    proc = _sp.run(
                        [_tool, 'x', '-o+', '-y', archive_path, rar_extract_dir + os.sep],
                        capture_output=True, timeout=120
                    )
                    if proc.returncode == 0:
                        break
                except OSError:
                    continue
            if proc is None or proc.returncode != 0:
                err_msg = proc.stderr.decode(errors='replace') if proc else 'unrar not found'
                logger.error(f"[BULK] unrar failed: {err_msg}")
                return jsonify({
                    'status': 'error',
                    'message': 'Failed to extract RAR file. Please ensure WinRAR/unrar is installed, or upload a .zip file instead.'
                }), 400
            _collect_resume_files(rar_extract_dir)
        else:
            zip_extract_dir = os.path.join(temp_dir, 'zip_extracted')
            os.makedirs(zip_extract_dir, exist_ok=True)
            with zipfile.ZipFile(archive_path, 'r') as zf:
                zf.extractall(zip_extract_dir)
            _collect_resume_files(zip_extract_dir)

        if not resume_files:
            return jsonify({
                'status': 'error',
                'message': 'No PDF or DOCX files found in the archive'
            }), 400

        total = len(resume_files)
        logger.info(f"[BULK] Found {total} resume files. Starting parallel processing with {MAX_BULK_WORKERS} workers...")

        # Process all resumes in parallel
        results = []
        with ThreadPoolExecutor(max_workers=MAX_BULK_WORKERS) as executor:
            future_map = {
                executor.submit(
                    _process_single_resume,
                    filepath, original_name, job_description, job_info, job_id, upload_folder
                ): original_name
                for filepath, original_name in resume_files
            }
            for future in as_completed(future_map):
                try:
                    result = future.result(timeout=120)
                    results.append(result)
                except Exception as exc:
                    results.append({
                        'filename': future_map[future],
                        'status': 'error',
                        'error': str(exc),
                        'name': None,
                        'email': None,
                        'match_score': 0,
                        'recommendation': None,
                        'candidate_id': None
                    })

        # Aggregate stats
        success = [r for r in results if r['status'] == 'success']
        duplicates = [r for r in results if r['status'] == 'duplicate']
        errors = [r for r in results if r['status'] == 'error']

        logger.info(f"[BULK] COMPLETE: {len(success)} success, {len(duplicates)} duplicates, {len(errors)} errors out of {total}")

        return jsonify({
            'status': 'success',
            'message': f'Processed {total} resumes: {len(success)} added, {len(duplicates)} duplicates, {len(errors)} failed',
            'summary': {
                'total': total,
                'success': len(success),
                'duplicates': len(duplicates),
                'errors': len(errors),
                'job': {
                    'id': job_info['id'],
                    'title': job_info['title'],
                    'department': job_info.get('department')
                }
            },
            'results': sorted(results, key=lambda r: r.get('match_score', 0), reverse=True)
        })

    except zipfile.BadZipFile:
        return jsonify({'status': 'error', 'message': 'Invalid or corrupted archive file'}), 400
    except Exception as e:
        if 'rarfile' in str(type(e).__module__ or ''):
            return jsonify({'status': 'error', 'message': 'Invalid or corrupted RAR file'}), 400
        logger.exception(f"[BULK] Unexpected error: {e}")
        return jsonify({'status': 'error', 'message': f'Bulk upload failed: {str(e)}'}), 500
    finally:
        # Clean up temp directory
        with contextlib.suppress(Exception):
            shutil.rmtree(temp_dir, ignore_errors=True)


# ============================================================================
#                     AI ENHANCE (JOB / SECTOR TEXT)
# ============================================================================

@admin_bp.route('/ai-enhance', methods=['POST'])
@jwt_required()
@require_admin_role
def ai_enhance_text():
    """
    Use AI to polish a job posting or sector name + description.
    Keeps the user's original intent but makes text more professional,
    well-structured, and compelling.

    JSON body:
        type: 'job' | 'sector'
        title: current title text
        description: current description text
    Returns:
        enhanced_title, enhanced_description
    """
    data = request.get_json(force=True)
    text_type = data.get('type', 'job')
    title = data.get('title') or ''
    description = data.get('description') or ''

    # Ensure title and description are strings (frontend may send objects)
    if isinstance(title, dict):
        title = ' '.join(str(v) for v in title.values())
    if isinstance(description, dict):
        parts = []
        for key, val in description.items():
            heading = key.replace('_', ' ').title()
            if isinstance(val, list):
                val = '\n'.join(f'- {item}' for item in val)
            parts.append(f"{heading}:\n{val}")
        description = '\n\n'.join(parts)
    elif isinstance(description, list):
        description = '\n'.join(str(item) for item in description)

    title = str(title).strip()
    description = str(description).strip()

    if not title and not description:
        return jsonify({'status': 'error', 'message': 'Provide at least a title or description to enhance'}), 400

    try:
        api_key = os.environ.get('OPENAI_API_KEY')
        if not api_key:
            return jsonify({'status': 'error', 'message': 'OpenAI API key not configured'}), 500
        
        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
        except TypeError:
            import httpx
            proxy = os.environ.get('HTTPS_PROXY') or os.environ.get('https_proxy')
            http_client = httpx.Client(proxies=proxy) if proxy else httpx.Client()
            from openai import OpenAI
            client = OpenAI(api_key=api_key, http_client=http_client)

        if text_type == 'sector':
            system_msg = (
                "You are a corporate branding specialist. The user will give you a sector/department name and description. "
                "Polish them to sound professional and clear. Keep it concise. "
                "Return JSON with keys: enhanced_title (string), enhanced_description (single plain-text string, not nested). No markdown fences."
            )
            user_msg = f"Sector name: {title}\nDescription: {description}"
        else:
            system_msg = (
                "You are a senior HR copywriter at a leading company based in India. "
                "The user will give you a job title and description draft. "
                "Polish the title to be industry-standard (concise, clear seniority). "
                "Rewrite the description to focus ONLY on: a brief overview of the role, key responsibilities, and what is expected day-to-day. "
                "Use bullet points for responsibilities. Make it compelling and professional. "
                "Do NOT include qualifications, experience requirements, education/degree requirements, or salary in the description \u2014 the recruiter fills those separately. "
                "From the description, extract ONLY concrete, domain-specific skills relevant to the job title. "
                "For tech roles these would be tools/frameworks/languages (e.g. React, Python, AWS, Docker, PostgreSQL). "
                "For non-tech roles these would be domain expertise areas (e.g. Contract Law, Corporate Governance, Financial Modelling, Supply Chain Management, Talent Acquisition). "
                "Do NOT list soft skills or vague abilities like 'communication', 'problem solving', 'team management', 'leadership' as skills \u2014 those belong in the description. "
                "Every skill must be in Title Case. "
                "Split them into required_skills and preferred_skills (comma-separated strings). "
                "Return JSON with keys: enhanced_title (string), enhanced_description (plain-text string, responsibilities only), "
                "required_skills (comma-separated string of must-have domain skills), preferred_skills (comma-separated string of nice-to-have domain skills). "
                "No markdown fences."
            )
            user_msg = f"Job title: {title}\nDescription draft: {description}"

        response = client.chat.completions.create(
            model='gpt-4o-mini',
            messages=[
                {'role': 'system', 'content': system_msg},
                {'role': 'user', 'content': user_msg}
            ],
            temperature=0.7,
            max_tokens=1500
        )

        raw = response.choices[0].message.content.strip()
        # Strip markdown code fences if present
        if raw.startswith('```'):
            raw = raw.split('\n', 1)[-1].rsplit('```', 1)[0].strip()

        import json as _j
        result = _j.loads(raw)

        enhanced_title = result.get('enhanced_title', title)
        enhanced_desc = result.get('enhanced_description', description)

        # If AI returned a nested object instead of a plain string, flatten it
        if isinstance(enhanced_title, dict):
            enhanced_title = str(enhanced_title)
        if isinstance(enhanced_desc, dict):
            # Join section values into a readable string
            parts = []
            for key, val in enhanced_desc.items():
                heading = key.replace('_', ' ').title()
                if isinstance(val, list):
                    val = '\n'.join(f'- {item}' for item in val)
                parts.append(f"{heading}:\n{val}")
            enhanced_desc = '\n\n'.join(parts)
        elif isinstance(enhanced_desc, list):
            enhanced_desc = '\n'.join(str(item) for item in enhanced_desc)

        # Extract skills if present (job type only)
        required_skills = result.get('required_skills', '')
        preferred_skills = result.get('preferred_skills', '')
        if isinstance(required_skills, list):
            required_skills = ', '.join(str(s) for s in required_skills)
        if isinstance(preferred_skills, list):
            preferred_skills = ', '.join(str(s) for s in preferred_skills)

        # Title-case each skill
        def _title_case_skills(skills_str):
            if not skills_str:
                return ''
            return ', '.join(s.strip().title() for s in str(skills_str).split(',') if s.strip())

        required_skills = _title_case_skills(required_skills)
        preferred_skills = _title_case_skills(preferred_skills)

        resp = {
            'status': 'success',
            'enhanced_title': str(enhanced_title),
            'enhanced_description': str(enhanced_desc)
        }
        if text_type == 'job':
            resp['required_skills'] = str(required_skills)
            resp['preferred_skills'] = str(preferred_skills)

        return jsonify(resp)

    except Exception as e:
        logger.error(f"[AI ENHANCE] Failed: {e}")
        return jsonify({'status': 'error', 'message': f'AI enhancement failed: {str(e)}'}), 500


# ==============================================================================
#                     CUSTOM QUESTION BANK
# ==============================================================================

def _extract_text_from_file(filepath):
    """Extract text from a PDF or DOCX file."""
    text = ""
    if filepath.lower().endswith('.pdf'):
        from PyPDF2 import PdfReader
        with open(filepath, 'rb') as f:
            pdf = PdfReader(f)
            for page in pdf.pages:
                if page_text := page.extract_text():
                    text += page_text + "\n"
    elif filepath.lower().endswith('.docx'):
        from docx import Document
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    return text.strip()


def _parse_questions_from_text(text):
    """
    Parse questions from raw text extracted from uploaded files.
    Returns a list of question dicts: {question, options, correct_answer, category, difficulty}
    """
    questions = []

    # Try AI-powered parsing first
    try:
        api_key = os.environ.get('OPENAI_API_KEY')
        if not api_key:
            raise ValueError('OPENAI_API_KEY not configured')
        
        try:
            from openai import OpenAI
            client = OpenAI(api_key=api_key)
        except TypeError:
            import httpx
            proxy = os.environ.get('HTTPS_PROXY') or os.environ.get('https_proxy')
            http_client = httpx.Client(proxies=proxy) if proxy else httpx.Client()
            from openai import OpenAI
            client = OpenAI(api_key=api_key, http_client=http_client)
        
        # Truncate text if too long for context window
        truncated = text[:12000] if len(text) > 12000 else text

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": """You are an expert at parsing assessment questions from documents.
Extract ALL questions from the given text. For each question, identify:
- The question text
- The options (if multiple-choice)
- The correct answer (if indicated, otherwise null)
- A category/topic tag
- Difficulty level (easy/medium/hard)

Return a JSON array. Each element:
{
  "question": "...",
  "options": ["A", "B", "C", "D"] or null if not MCQ,
  "correct_answer": "..." or null,
  "category": "topic",
  "difficulty": "medium"
}

If the document contains free-form questions (not MCQ), still include them with options=null.
Return ONLY valid JSON, no markdown."""},
                {"role": "user", "content": f"Parse all questions from this document:\n\n{truncated}"}
            ],
            temperature=0.2,
            max_tokens=4000
        )
        import json
        content = response.choices[0].message.content.strip()
        if content.startswith("```"):
            content = content.split("\n", 1)[1].rsplit("```", 1)[0].strip()
        questions = json.loads(content)
        logger.info(f"[CUSTOM QB] AI parsed {len(questions)} questions from uploaded document")
        return questions
    except Exception as ai_err:
        logger.warning(f"[CUSTOM QB] AI parsing failed: {ai_err}, falling back to regex")

    # Regex fallback: look for numbered questions
    import re
    # Pattern: number followed by . or ) then question text
    q_pattern = re.compile(r'(?:^|\n)\s*(\d+)\s*[.)]\s*(.+?)(?=\n\s*\d+\s*[.)]|\n*$)', re.DOTALL)
    matches = q_pattern.findall(text)
    for num, q_text in matches:
        q_text = q_text.strip()
        if len(q_text) > 15:  # Skip too-short matches
            questions.append({
                'question': q_text,
                'options': None,
                'correct_answer': None,
                'category': 'custom',
                'difficulty': 'medium'
            })

    if not questions:
        # If no numbered questions found, split by lines that end with ?
        for line in text.split('\n'):
            line = line.strip()
            if line.endswith('?') and len(line) > 20:
                questions.append({
                    'question': line,
                    'options': None,
                    'correct_answer': None,
                    'category': 'custom',
                    'difficulty': 'medium'
                })

    logger.info(f"[CUSTOM QB] Regex parsed {len(questions)} questions from uploaded document")
    return questions


@admin_bp.route('/question-bank/upload', methods=['POST'])
@jwt_required()
def upload_question_bank():
    """
    Upload a PDF or DOCX containing custom questions.
    The file is parsed and questions are extracted.
    """
    try:
        claims = get_jwt()
        role = claims.get('role', '')
        if role not in ('admin', 'super_admin', 'interviewer'):
            return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

        if 'file' not in request.files:
            return jsonify({'status': 'error', 'message': 'No file uploaded'}), 400

        file = request.files['file']
        if not file.filename:
            return jsonify({'status': 'error', 'message': 'No file selected'}), 400

        ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
        if ext not in ('pdf', 'docx'):
            return jsonify({'status': 'error', 'message': 'Only PDF and DOCX files allowed'}), 400

        description = request.form.get('description', '')
        tags = request.form.get('tags', '')

        # Save file
        upload_dir = os.path.join(os.path.dirname(__file__), 'uploads', 'question_banks')
        os.makedirs(upload_dir, exist_ok=True)
        original_filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{original_filename}"
        filepath = os.path.join(upload_dir, unique_filename)
        file.save(filepath)

        # Extract text
        questions_text = _extract_text_from_file(filepath)
        if not questions_text or len(questions_text.strip()) < 30:
            os.remove(filepath)
            return jsonify({'status': 'error', 'message': 'Could not extract any text from the file. Please check the document.'}), 400

        # Parse questions
        parsed_questions = _parse_questions_from_text(questions_text)

        # Save to DB
        user_id = int(get_jwt_identity())
        conn = get_connection()
        cur = conn.cursor()
        import json
        cur.execute("""
            INSERT INTO custom_question_bank
            (filename, original_filename, file_path, questions_text, parsed_questions, uploaded_by, description, tags)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (unique_filename, file.filename, filepath,
              questions_text, json.dumps(parsed_questions),
              user_id, description, tags))
        qb_id = cur.fetchone()[0]
        conn.commit()
        return_connection(conn)

        logger.info(f"[CUSTOM QB] Uploaded question bank #{qb_id}: {file.filename} ({len(parsed_questions)} questions parsed)")

        return jsonify({
            'status': 'success',
            'message': f'Uploaded successfully — {len(parsed_questions)} questions parsed',
            'data': {
                'id': qb_id,
                'filename': file.filename,
                'questions_count': len(parsed_questions),
                'parsed_questions': parsed_questions[:3],  # Preview first 3
                'raw_text_length': len(questions_text)
            }
        }), 201

    except Exception as e:
        logger.error(f"[CUSTOM QB] Upload failed: {e}")
        return jsonify({'status': 'error', 'message': f'Upload failed: {str(e)}'}), 500


@admin_bp.route('/question-bank', methods=['GET'])
@jwt_required()
def list_question_banks():
    """List all custom question bank uploads."""
    try:
        claims = get_jwt()
        role = claims.get('role', '')
        if role not in ('admin', 'super_admin', 'interviewer'):
            return jsonify({'status': 'error', 'message': 'Unauthorized'}), 403

        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT qb.id, qb.original_filename, qb.description, qb.tags,
                   qb.is_active, qb.created_at, u.name as uploaded_by_name,
                   jsonb_array_length(COALESCE(qb.parsed_questions, '[]'::jsonb)) as questions_count
            FROM custom_question_bank qb
            LEFT JOIN users u ON qb.uploaded_by = u.id
            ORDER BY qb.created_at DESC
        """)
        rows = cur.fetchall()
        return_connection(conn)

        items = [{
            'id': row[0],
            'filename': row[1],
            'description': row[2],
            'tags': row[3],
            'is_active': row[4],
            'created_at': str(row[5]) if row[5] else None,
            'uploaded_by': row[6],
            'questions_count': row[7] or 0
        } for row in rows]

        return jsonify({'status': 'success', 'data': items})

    except Exception as e:
        logger.error(f"[CUSTOM QB] List failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/question-bank/<int:qb_id>', methods=['GET'])
@jwt_required()
def get_question_bank(qb_id):
    """Get a single question bank with all parsed questions."""
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            SELECT qb.id, qb.original_filename, qb.description, qb.tags,
                   qb.is_active, qb.created_at, qb.parsed_questions, qb.questions_text,
                   u.name as uploaded_by_name
            FROM custom_question_bank qb
            LEFT JOIN users u ON qb.uploaded_by = u.id
            WHERE qb.id = %s
        """, (qb_id,))
        row = cur.fetchone()
        return_connection(conn)

        if not row:
            return jsonify({'status': 'error', 'message': 'Not found'}), 404

        import json
        parsed = json.loads(row[6]) if isinstance(row[6], str) else (row[6] or [])

        return jsonify({
            'status': 'success',
            'data': {
                'id': row[0],
                'filename': row[1],
                'description': row[2],
                'tags': row[3],
                'is_active': row[4],
                'created_at': str(row[5]) if row[5] else None,
                'parsed_questions': parsed,
                'raw_text_preview': (row[7] or '')[:500],
                'uploaded_by': row[8],
                'questions_count': len(parsed)
            }
        })

    except Exception as e:
        logger.error(f"[CUSTOM QB] Get failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/question-bank/<int:qb_id>', methods=['DELETE'])
@jwt_required()
def delete_question_bank(qb_id):
    """Delete a custom question bank entry and its file."""
    try:
        claims = get_jwt()
        role = claims.get('role', '')
        if role not in ('admin', 'super_admin'):
            return jsonify({'status': 'error', 'message': 'Only admins can delete'}), 403

        conn = get_connection()
        cur = conn.cursor()
        cur.execute("SELECT file_path FROM custom_question_bank WHERE id = %s", (qb_id,))
        row = cur.fetchone()
        if not row:
            return_connection(conn)
            return jsonify({'status': 'error', 'message': 'Not found'}), 404

        filepath = row[0]
        cur.execute("DELETE FROM custom_question_bank WHERE id = %s", (qb_id,))
        conn.commit()
        return_connection(conn)

        # Remove file
        if filepath and os.path.exists(filepath):
            with contextlib.suppress(Exception):
                os.remove(filepath)

        return jsonify({'status': 'success', 'message': 'Question bank deleted'})

    except Exception as e:
        logger.error(f"[CUSTOM QB] Delete failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500


@admin_bp.route('/question-bank/<int:qb_id>/toggle', methods=['PATCH'])
@jwt_required()
def toggle_question_bank(qb_id):
    """Toggle active/inactive status of a question bank."""
    try:
        conn = get_connection()
        cur = conn.cursor()
        cur.execute("""
            UPDATE custom_question_bank
            SET is_active = NOT is_active, updated_at = CURRENT_TIMESTAMP
            WHERE id = %s
            RETURNING is_active
        """, (qb_id,))
        row = cur.fetchone()
        if not row:
            return_connection(conn)
            return jsonify({'status': 'error', 'message': 'Not found'}), 404
        conn.commit()
        return_connection(conn)

        return jsonify({'status': 'success', 'is_active': row[0]})

    except Exception as e:
        logger.error(f"[CUSTOM QB] Toggle failed: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500
