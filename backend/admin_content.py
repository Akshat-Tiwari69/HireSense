"""
Admin content routes — bulk resume upload, AI text enhancement, and custom question banks.
"""

import os
import re
import uuid
import json as _json
import zipfile
import tempfile
import shutil
import contextlib
import logging
import stat
from pathlib import PurePosixPath
from concurrent.futures import ThreadPoolExecutor, as_completed
from flask import Blueprint, request, jsonify
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename
from db_config import db_connection, get_connection, return_connection
from candidate_db import get_candidate_by_email, insert_candidate_application
from admin_middleware import require_admin_role
from storage_config import get_upload_root, get_upload_subdirectory

logger = logging.getLogger(__name__)

admin_content_bp = Blueprint('admin_content', __name__)

ALLOWED_RESUME_EXTENSIONS = {'pdf', 'docx'}
EMAIL_PATTERN_BULK = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9]([a-zA-Z0-9-]*[a-zA-Z0-9])?(\.[a-zA-Z0-9]([a-zA-Z0-9-]*[a-zA-Z0-9])?)*\.[a-zA-Z]{2,}$'
MAX_BULK_WORKERS = 8
MAX_ARCHIVE_MEMBERS = 100
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_RESUME_BYTES = 10 * 1024 * 1024
MAX_QUESTION_TEXT_CHARS = 100_000
MAX_AI_INPUT_CHARS = 20_000


# ============================================================================
#                        BULK UPLOAD HELPERS
# ============================================================================

def _bulk_valid_email(email):
    if email is None or not isinstance(email, str) or not email.strip():
        return False
    return re.match(EMAIL_PATTERN_BULK, email) is not None


def _bulk_name_from_email(email):
    if not email or '@' not in email:
        return None
    local = email.split('@', 1)[0]
    parts = re.split(r'[._-]+', local)
    parts = [p for p in parts if p]
    return " ".join(p.capitalize() for p in parts) if parts else None


def _merge_ai_data_to_parsed(parsed_data, ai_data):
    if ai_data.get('skills'):
        parsed_data['skills'] = ai_data['skills']
    if ai_data.get('experience') and ai_data['experience'] > 0:
        parsed_data['experience'] = ai_data['experience']
    if ai_data.get('education'):
        parsed_data['education'] = ai_data['education']
    if ai_data.get('name'):
        parsed_data['name'] = ai_data['name']
    if ai_data.get('email'):
        parsed_data['email'] = ai_data['email']
    if ai_data.get('phone'):
        parsed_data['phone'] = ai_data['phone']


def _safe_zip_member_name(member):
    """Return a normalized archive name or reject traversal/symlink entries."""
    normalized = member.filename.replace('\\', '/')
    path = PurePosixPath(normalized)
    mode = member.external_attr >> 16
    if (
        path.is_absolute()
        or not path.parts
        or '..' in path.parts
        or ':' in path.parts[0]
        or stat.S_ISLNK(mode)
    ):
        raise ValueError('Archive contains an unsafe file path')
    if member.flag_bits & 0x1:
        raise ValueError('Encrypted archives are not supported')
    return path


def _extract_resume_zip(archive_path, destination):
    """Safely extract bounded PDF/DOCX members and return their display names."""
    extracted = []
    with zipfile.ZipFile(archive_path, 'r') as archive:
        members = archive.infolist()
        if len(members) > MAX_ARCHIVE_MEMBERS:
            raise ValueError(f'Archive contains more than {MAX_ARCHIVE_MEMBERS} files')

        total_size = sum(member.file_size for member in members if not member.is_dir())
        if total_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
            raise ValueError('Archive expands beyond the allowed size')

        for member in members:
            path = _safe_zip_member_name(member)
            if member.is_dir() or '__MACOSX' in path.parts or path.name.startswith('.'):
                continue
            extension = path.suffix.lower().lstrip('.')
            if extension not in ALLOWED_RESUME_EXTENSIONS:
                continue
            if member.file_size > MAX_RESUME_BYTES:
                raise ValueError(f'Resume {path.name} exceeds the 10 MB limit')
            if member.compress_size == 0 and member.file_size:
                raise ValueError('Archive contains an invalid compressed entry')
            if member.compress_size and member.file_size / member.compress_size > 200:
                raise ValueError('Archive contains a suspiciously compressed entry')

            safe_name = secure_filename(path.name)
            if not safe_name:
                raise ValueError('Archive contains a resume with an invalid filename')
            extracted_path = os.path.join(destination, f'{uuid.uuid4()}_{safe_name}')
            bytes_written = 0
            with archive.open(member, 'r') as source, open(extracted_path, 'wb') as target:
                while chunk := source.read(1024 * 1024):
                    bytes_written += len(chunk)
                    if bytes_written > MAX_RESUME_BYTES:
                        raise ValueError(f'Resume {path.name} exceeds the 10 MB limit')
                    target.write(chunk)
            extracted.append((extracted_path, str(path)))
    return extracted


def _create_openai_client():
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        raise RuntimeError('OPENAI_API_KEY not configured')
    from openai import OpenAI
    return OpenAI(api_key=api_key, timeout=30.0, max_retries=1)


def _parse_ai_json(raw):
    if not isinstance(raw, str) or not raw.strip():
        raise ValueError('AI provider returned an empty response')
    content = raw.strip()
    if content.startswith('```'):
        content = content.split('\n', 1)[-1].rsplit('```', 1)[0].strip()
    return _json.loads(content)


def _normalize_parsed_questions(value):
    if not isinstance(value, list):
        raise ValueError('Question parser did not return a list')
    normalized = []
    for item in value[:200]:
        if not isinstance(item, dict):
            continue
        question = item.get('question')
        if not isinstance(question, str) or not question.strip():
            continue
        options = item.get('options')
        if options is not None:
            if not isinstance(options, list):
                options = None
            else:
                options = [str(option).strip()[:500] for option in options[:20] if str(option).strip()]
        correct_answer = item.get('correct_answer')
        normalized.append({
            'question': question.strip()[:2_000],
            'options': options,
            'correct_answer': str(correct_answer).strip()[:500] if correct_answer is not None else None,
            'category': str(item.get('category') or 'custom').strip()[:100],
            'difficulty': str(item.get('difficulty') or 'medium').lower()
            if str(item.get('difficulty') or 'medium').lower() in {'easy', 'medium', 'hard'}
            else 'medium',
        })
    return normalized


def _fetch_job_for_bulk(job_id):
    conn = None
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, title, department, required_skills, preferred_skills, min_experience "
            "FROM job_descriptions WHERE id = %s AND status = 'active'",
            (int(job_id),)
        )
        row = cursor.fetchone()
        if row:
            skills = set()
            for skills_val in (row[3], row[4]):
                if skills_val:
                    with contextlib.suppress(ValueError, TypeError):
                        parsed = _json.loads(skills_val)
                        if isinstance(parsed, list):
                            skills.update(s.strip() for s in parsed if s.strip())
                            continue
                    skills.update(s.strip() for s in str(skills_val).split(',') if s.strip())
            min_exp = row[5] or 0
            job_info = {'id': row[0], 'title': row[1], 'department': row[2]}
            return {'skills': list(skills), 'min_experience': min_exp, 'title': row[1], 'department': row[2]}, job_info
    except Exception as e:
        logger.warning(f"[BULK] Could not load job posting {job_id}: {e}")
    finally:
        if conn:
            return_connection(conn)
    return None, None


def _process_single_resume(filepath, filename, job_description, job_info, job_id):
    from resume_parser import parse_resume, calculate_match_score
    from resume_analyzer import analyze_resume, ResumeAnalyzer

    result = {
        'filename': filename, 'status': 'error',
        'name': None, 'email': None, 'match_score': 0,
        'recommendation': None, 'candidate_id': None,
        'error': None, 'missing': []
    }

    try:
        parsed_data = parse_resume(filepath, job_description)

        with open(filepath, 'rb') as f:
            if filepath.lower().endswith('.pdf'):
                from PyPDF2 import PdfReader
                pdf = PdfReader(f)
                resume_text = " ".join([page.extract_text() or '' for page in pdf.pages])
            else:
                from docx import Document
                doc = Document(f)
                resume_text = " ".join([para.text for para in doc.paragraphs])

        if not resume_text or len(resume_text.strip()) < 50:
            result['missing'] = ['name', 'email']
            name = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title() or 'Unknown Candidate'
            email = f'unknown-{uuid.uuid4().hex[:12]}@bulk-upload.local'
            try:
                candidate_id = insert_candidate_application(
                    name=name, email=email, phone='',
                    resume_path=filepath,
                    parsed_data={'skills': [], 'experience': 0, 'education': '', 'match_score': 0, 'shortlist_status': 'Pending Review'},
                    job_id=int(job_id),
                    ai_reasoning='Resume text could not be extracted; manual review required.',
                    pros=None, cons=None, status='absence_of_details',
                )
                result.update({'status': 'success', 'name': name, 'email': email, 'candidate_id': candidate_id,
                                'error': 'Could not extract text — saved with Absence of Details'})
            except Exception as save_err:
                result['error'] = f'Could not extract text and save failed: {save_err}'
            return result

        try:
            analyzer = ResumeAnalyzer()
            if ai_data := analyzer.extract_resume_data(resume_text):
                _merge_ai_data_to_parsed(parsed_data, ai_data)
                parsed_data['match_score'] = calculate_match_score(
                    parsed_data.get('skills', []), parsed_data.get('experience', 0),
                    job_description.get('skills', []), job_description.get('min_experience', 0)
                )
        except Exception as ai_err:
            logger.warning(f"[BULK] AI extraction failed for {filename}: {ai_err}")

        ai_analysis = None
        try:
            ai_analysis = analyze_resume(
                resume_text=resume_text, parsed_data=parsed_data,
                job_requirements=job_description, enhance_score=True
            )
            if ai_analysis and 'enhanced_match_score' in ai_analysis:
                parsed_data['match_score'] = ai_analysis['enhanced_match_score']
        except Exception as ai_err:
            logger.warning(f"[BULK] AI analysis failed for {filename}: {ai_err}")
            ai_analysis = {
                "pros": ["Resume uploaded successfully"],
                "cons": ["AI analysis unavailable - manual review recommended"],
                "overall_assessment": "AI analysis failed. Manual review required.",
                "recommendation": "Pending Review", "confidence_score": 0
            }

        name = parsed_data.get('name')
        email = parsed_data.get('email')
        phone = parsed_data.get('phone', '')

        missing_details = []
        if not name:
            missing_details.append('name')
        if not email or not _bulk_valid_email(email):
            missing_details.append('email')

        if not email or not _bulk_valid_email(email):
            email = f'unknown-{uuid.uuid4().hex[:12]}@bulk-upload.local'

        if not name:
            name = _bulk_name_from_email(email)
            if not name or 'unknown' in name.lower():
                name = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title() or 'Unknown Candidate'

        candidate_status = 'absence_of_details' if missing_details else 'applied'
        result['name'] = name
        result['email'] = email
        result['match_score'] = parsed_data.get('match_score', 0)
        result['recommendation'] = ai_analysis.get('recommendation', 'Pending Review') if ai_analysis else 'Pending Review'
        result['missing'] = missing_details

        if not email.endswith('@bulk-upload.local'):
            with contextlib.suppress(Exception):
                if existing := get_candidate_by_email(email):
                    result.update({'status': 'duplicate', 'candidate_id': existing['id'],
                                   'error': f'Already registered (ID: {existing["id"]})'})
                    with contextlib.suppress(OSError):
                        os.remove(filepath)
                    return result

        pros_text = "\n".join(ai_analysis.get('pros', [])) if ai_analysis else None
        cons_text = "\n".join(ai_analysis.get('cons', [])) if ai_analysis else None

        candidate_id = insert_candidate_application(
            name=name, email=email, phone=phone or '',
            resume_path=filepath, parsed_data=parsed_data,
            job_id=int(job_id),
            ai_reasoning=ai_analysis.get('overall_assessment', '') if ai_analysis else '',
            pros=pros_text, cons=cons_text, status=candidate_status,
        )
        result['candidate_id'] = candidate_id

        result['status'] = 'success'
        if missing_details:
            result['error'] = f'Saved with Absence of Details (missing: {", ".join(missing_details)})'
        logger.info(f"[BULK] Processed {filename} -> {name} <{email}> score={result['match_score']} status={candidate_status}")

    except Exception as e:
        result['error'] = 'Resume processing failed'
        logger.exception("[BULK] Error processing %s: %s", filename, e)
        with contextlib.suppress(OSError):
            os.remove(filepath)

    return result


# ============================================================================
#                        QUESTION BANK HELPERS
# ============================================================================

def _extract_text_from_file(filepath):
    chunks = []
    total_chars = 0

    def append_text(value):
        nonlocal total_chars
        if not value:
            return
        total_chars += len(value) + 1
        if total_chars > MAX_QUESTION_TEXT_CHARS:
            raise ValueError('Extracted question-bank text is too large')
        chunks.append(value)

    if filepath.lower().endswith('.pdf'):
        from PyPDF2 import PdfReader
        with open(filepath, 'rb') as f:
            pdf = PdfReader(f)
            for page in pdf.pages:
                if page_text := page.extract_text():
                    append_text(page_text)
    elif filepath.lower().endswith('.docx'):
        from docx import Document
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                append_text(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        append_text(cell.text)
    return "\n".join(chunks).strip()


def _parse_questions_from_text(text):
    questions = []

    try:
        client = _create_openai_client()

        truncated = text[:12000] if len(text) > 12000 else text

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": """You are an expert at parsing assessment questions from documents.
Extract ALL questions from the given text. For each question, identify:
- The question text
- The options (if multiple-choice)
- The correct answer (if indicated, otherwise null)
- A category/topic tag
- Difficulty level (easy/medium/hard)

Return a JSON array. Each element:
{
  "question": "...",
  "options": ["A", "B", "C", "D"] or null if not MCQ,
  "correct_answer": "..." or null,
  "category": "topic",
  "difficulty": "medium"
}

If the document contains free-form questions (not MCQ), still include them with options=null.
Return ONLY valid JSON, no markdown."""},
                {"role": "user", "content": f"Parse all questions from this document:\n\n{truncated}"}
            ],
            temperature=0.2,
            max_tokens=4000
        )
        questions = _normalize_parsed_questions(
            _parse_ai_json(response.choices[0].message.content)
        )
        logger.info(f"[CUSTOM QB] AI parsed {len(questions)} questions from uploaded document")
        return questions
    except Exception as ai_err:
        logger.warning(f"[CUSTOM QB] AI parsing failed: {ai_err}, falling back to regex")

    q_pattern = re.compile(r'(?:^|\n)\s*(\d+)\s*[.)]\s*(.+?)(?=\n\s*\d+\s*[.)]|\n*$)', re.DOTALL)
    matches = q_pattern.findall(text)
    for num, q_text in matches:
        q_text = q_text.strip()
        if len(q_text) > 15:
            questions.append({'question': q_text, 'options': None, 'correct_answer': None,
                              'category': 'custom', 'difficulty': 'medium'})

    if not questions:
        for line in text.split('\n'):
            line = line.strip()
            if line.endswith('?') and len(line) > 20:
                questions.append({'question': line, 'options': None, 'correct_answer': None,
                                  'category': 'custom', 'difficulty': 'medium'})

    logger.info(f"[CUSTOM QB] Regex parsed {len(questions)} questions from uploaded document")
    return questions


# ============================================================================
#                        ROUTES
# ============================================================================

@admin_content_bp.route('/bulk-upload', methods=['POST'])
@jwt_required()
@require_admin_role
def bulk_upload_resumes():
    logger.info("[BULK] BULK RESUME UPLOAD REQUEST RECEIVED")

    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': 'No file uploaded'}), 400

    archive_file = request.files['file']
    fname = (archive_file.filename or '').lower()
    if not fname.endswith('.zip'):
        return jsonify({'status': 'error', 'message': 'Please upload a .zip file'}), 400

    job_id = request.form.get('job_id')
    if not job_id:
        return jsonify({'status': 'error', 'message': 'Please select a job position'}), 400

    job_description, job_info = _fetch_job_for_bulk(job_id)
    if not job_description:
        return jsonify({'status': 'error', 'message': 'Selected job is no longer active'}), 400

    logger.info(f"[BULK] Target job: {job_info['title']} (ID: {job_id})")

    temp_dir = tempfile.mkdtemp(prefix='bulk_upload_')
    upload_folder = str(get_upload_root(create=True))

    try:
        archive_path = os.path.join(temp_dir, 'upload.zip')
        archive_file.save(archive_path)
        extract_dir = os.path.join(temp_dir, 'resumes')
        os.makedirs(extract_dir, exist_ok=True)
        extracted_files = _extract_resume_zip(archive_path, extract_dir)
        resume_files = []
        for extracted_path, display_name in extracted_files:
            permanent_path = os.path.join(upload_folder, os.path.basename(extracted_path))
            shutil.copy2(extracted_path, permanent_path)
            resume_files.append((permanent_path, display_name))

        if not resume_files:
            return jsonify({'status': 'error', 'message': 'No PDF or DOCX files found in the archive'}), 400

        total = len(resume_files)
        logger.info(f"[BULK] Found {total} resume files. Starting parallel processing with {MAX_BULK_WORKERS} workers...")

        results = []
        with ThreadPoolExecutor(max_workers=MAX_BULK_WORKERS) as executor:
            future_map = {
                executor.submit(_process_single_resume, filepath, original_name,
                                job_description, job_info, job_id): (original_name, filepath)
                for filepath, original_name in resume_files
            }
            for future in as_completed(future_map):
                try:
                    results.append(future.result())
                except Exception:
                    original_name, filepath = future_map[future]
                    with contextlib.suppress(OSError):
                        os.remove(filepath)
                    logger.exception("[BULK] Worker failed for %s", original_name)
                    results.append({
                        'filename': original_name, 'status': 'error', 'error': 'Resume processing failed',
                        'name': None, 'email': None, 'match_score': 0,
                        'recommendation': None, 'candidate_id': None
                    })

        success = [r for r in results if r['status'] == 'success']
        duplicates = [r for r in results if r['status'] == 'duplicate']
        errors = [r for r in results if r['status'] == 'error']

        logger.info(f"[BULK] COMPLETE: {len(success)} success, {len(duplicates)} duplicates, {len(errors)} errors out of {total}")

        return jsonify({
            'status': 'success',
            'message': f'Processed {total} resumes: {len(success)} added, {len(duplicates)} duplicates, {len(errors)} failed',
            'summary': {
                'total': total, 'success': len(success), 'duplicates': len(duplicates), 'errors': len(errors),
                'job': {'id': job_info['id'], 'title': job_info['title'], 'department': job_info.get('department')}
            },
            'results': sorted(results, key=lambda r: r.get('match_score', 0), reverse=True)
        })

    except zipfile.BadZipFile:
        return jsonify({'status': 'error', 'message': 'Invalid or corrupted archive file'}), 400
    except ValueError as error:
        return jsonify({'status': 'error', 'message': str(error)}), 400
    except Exception as e:
        logger.exception(f"[BULK] Unexpected error: {e}")
        return jsonify({'status': 'error', 'message': 'Internal server error'}), 500
    finally:
        with contextlib.suppress(Exception):
            shutil.rmtree(temp_dir, ignore_errors=True)


@admin_content_bp.route('/ai-enhance', methods=['POST'])
@jwt_required()
@require_admin_role
def ai_enhance_text():
    data = request.get_json(silent=True)
    if not isinstance(data, dict):
        return jsonify({'status': 'error', 'message': 'A JSON object is required'}), 400
    text_type = data.get('type', 'job')
    if text_type not in {'job', 'sector'}:
        return jsonify({'status': 'error', 'message': 'Type must be job or sector'}), 400
    title = data.get('title') or ''
    description = data.get('description') or ''

    if isinstance(title, dict):
        title = ' '.join(str(v) for v in title.values())
    if isinstance(description, dict):
        parts = []
        for key, val in description.items():
            heading = key.replace('_', ' ').title()
            if isinstance(val, list):
                val = '\n'.join(f'- {item}' for item in val)
            parts.append(f"{heading}:\n{val}")
        description = '\n\n'.join(parts)
    elif isinstance(description, list):
        description = '\n'.join(str(item) for item in description)

    title = str(title).strip()
    description = str(description).strip()

    if not title and not description:
        return jsonify({'status': 'error', 'message': 'Provide at least a title or description to enhance'}), 400
    if len(title) > 300 or len(description) > MAX_AI_INPUT_CHARS:
        return jsonify({'status': 'error', 'message': 'Input is too long'}), 400

    try:
        client = _create_openai_client()

        if text_type == 'sector':
            system_msg = (
                "You are a corporate branding specialist. The user will give you a sector/department name and description. "
                "Polish them to sound professional and clear. Keep it concise. "
                "Return JSON with keys: enhanced_title (string), enhanced_description (single plain-text string, not nested). No markdown fences."
            )
            user_msg = f"Sector name: {title}\nDescription: {description}"
        else:
            system_msg = (
                "You are a senior HR copywriter at a leading company based in India. "
                "The user will give you a job title and description draft. "
                "Polish the title to be industry-standard (concise, clear seniority). "
                "Rewrite the description to focus ONLY on: a brief overview of the role, key responsibilities, and what is expected day-to-day. "
                "Use bullet points for responsibilities. Make it compelling and professional. "
                "Do NOT include qualifications, experience requirements, education/degree requirements, or salary in the description — the recruiter fills those separately. "
                "From the description, extract ONLY concrete, domain-specific skills relevant to the job title. "
                "For tech roles these would be tools/frameworks/languages (e.g. React, Python, AWS, Docker, PostgreSQL). "
                "For non-tech roles these would be domain expertise areas (e.g. Contract Law, Corporate Governance, Financial Modelling, Supply Chain Management, Talent Acquisition). "
                "Do NOT list soft skills or vague abilities like 'communication', 'problem solving', 'team management', 'leadership' as skills — those belong in the description. "
                "Every skill must be in Title Case. "
                "Split them into required_skills and preferred_skills (comma-separated strings). "
                "Return JSON with keys: enhanced_title (string), enhanced_description (plain-text string, responsibilities only), "
                "required_skills (comma-separated string of must-have domain skills), preferred_skills (comma-separated string of nice-to-have domain skills). "
                "No markdown fences."
            )
            user_msg = f"Job title: {title}\nDescription draft: {description}"

        response = client.chat.completions.create(
            model='gpt-4o-mini',
            messages=[
                {'role': 'system', 'content': system_msg},
                {'role': 'user', 'content': user_msg}
            ],
            temperature=0.7,
            max_tokens=1500
        )

        result = _parse_ai_json(response.choices[0].message.content)
        if not isinstance(result, dict):
            raise ValueError('AI provider returned an invalid response')

        enhanced_title = result.get('enhanced_title', title)
        enhanced_desc = result.get('enhanced_description', description)

        if isinstance(enhanced_title, dict):
            enhanced_title = str(enhanced_title)
        if isinstance(enhanced_desc, dict):
            parts = []
            for key, val in enhanced_desc.items():
                heading = key.replace('_', ' ').title()
                if isinstance(val, list):
                    val = '\n'.join(f'- {item}' for item in val)
                parts.append(f"{heading}:\n{val}")
            enhanced_desc = '\n\n'.join(parts)
        elif isinstance(enhanced_desc, list):
            enhanced_desc = '\n'.join(str(item) for item in enhanced_desc)

        required_skills = result.get('required_skills', '')
        preferred_skills = result.get('preferred_skills', '')
        if isinstance(required_skills, list):
            required_skills = ', '.join(str(s) for s in required_skills)
        if isinstance(preferred_skills, list):
            preferred_skills = ', '.join(str(s) for s in preferred_skills)

        def _title_case_skills(skills_str):
            if not skills_str:
                return ''
            return ', '.join(s.strip().title() for s in str(skills_str).split(',') if s.strip())

        resp = {
            'status': 'success',
            'enhanced_title': str(enhanced_title),
            'enhanced_description': str(enhanced_desc)
        }
        if text_type == 'job':
            resp['required_skills'] = _title_case_skills(required_skills)
            resp['preferred_skills'] = _title_case_skills(preferred_skills)

        return jsonify(resp)

    except RuntimeError as error:
        logger.warning("[AI ENHANCE] Service unavailable: %s", error)
        return jsonify({'status': 'error', 'message': 'AI enhancement is not configured'}), 503
    except Exception:
        logger.exception("[AI ENHANCE] Provider request failed")
        return jsonify({'status': 'error', 'message': 'AI enhancement failed'}), 502


@admin_content_bp.route('/question-bank/upload', methods=['POST'])
@jwt_required()
@require_admin_role
def upload_question_bank():
    filepath = None
    stored = False
    try:
        if 'file' not in request.files:
            return jsonify({'status': 'error', 'message': 'No file uploaded'}), 400

        file = request.files['file']
        if not file.filename:
            return jsonify({'status': 'error', 'message': 'No file selected'}), 400

        ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
        if ext not in ('pdf', 'docx'):
            return jsonify({'status': 'error', 'message': 'Only PDF and DOCX files allowed'}), 400

        description = request.form.get('description', '')
        tags = request.form.get('tags', '')
        if len(description) > 2_000 or len(tags) > 500:
            return jsonify({'status': 'error', 'message': 'Description or tags are too long'}), 400

        upload_dir = str(get_upload_subdirectory('question_banks', create=True))
        original_filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{original_filename}"
        filepath = os.path.join(upload_dir, unique_filename)
        file.save(filepath)

        questions_text = _extract_text_from_file(filepath)
        if not questions_text or len(questions_text.strip()) < 30:
            os.remove(filepath)
            return jsonify({'status': 'error', 'message': 'Could not extract any text from the file.'}), 400

        parsed_questions = _parse_questions_from_text(questions_text)
        if not parsed_questions:
            return jsonify({'status': 'error', 'message': 'No valid questions were found'}), 400

        user_id = int(get_jwt_identity())
        with db_connection() as conn:
            cur = conn.cursor()
            cur.execute("""
                INSERT INTO custom_question_bank
                (filename, original_filename, file_path, questions_text, parsed_questions, uploaded_by, description, tags)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (unique_filename, original_filename, filepath,
                  questions_text, _json.dumps(parsed_questions),
                  user_id, description, tags))
            qb_id = cur.fetchone()[0]
            conn.commit()
        stored = True

        logger.info(f"[CUSTOM QB] Uploaded question bank #{qb_id}: {file.filename} ({len(parsed_questions)} questions parsed)")

        return jsonify({
            'status': 'success',
            'message': f'Uploaded successfully — {len(parsed_questions)} questions parsed',
            'data': {
                'id': qb_id, 'filename': file.filename,
                'questions_count': len(parsed_questions),
                'parsed_questions': parsed_questions[:3],
                'raw_text_length': len(questions_text)
            }
        }), 201

    except ValueError as error:
        return jsonify({'status': 'error', 'message': str(error)}), 400
    except Exception:
        logger.exception("[CUSTOM QB] Upload failed")
        return jsonify({'status': 'error', 'message': 'Internal server error'}), 500
    finally:
        if filepath and not stored:
            with contextlib.suppress(OSError):
                os.remove(filepath)


@admin_content_bp.route('/question-bank', methods=['GET'])
@jwt_required()
@require_admin_role
def list_question_banks():
    try:
        with db_connection() as conn:
            cur = conn.cursor()
            cur.execute("""
                SELECT qb.id, qb.original_filename, qb.description, qb.tags,
                       qb.is_active, qb.created_at, u.name as uploaded_by_name,
                       jsonb_array_length(COALESCE(qb.parsed_questions, '[]'::jsonb)) as questions_count
                FROM custom_question_bank qb
                LEFT JOIN users u ON qb.uploaded_by = u.id
                ORDER BY qb.created_at DESC
            """)
            rows = cur.fetchall()

        items = [{
            'id': row[0], 'filename': row[1], 'description': row[2],
            'tags': row[3], 'is_active': row[4],
            'created_at': str(row[5]) if row[5] else None,
            'uploaded_by': row[6], 'questions_count': row[7] or 0
        } for row in rows]

        return jsonify({'status': 'success', 'data': items})

    except Exception as e:
        logger.error(f"[CUSTOM QB] List failed: {e}")
        return jsonify({'status': 'error', 'message': 'Internal server error'}), 500


@admin_content_bp.route('/question-bank/<int:qb_id>', methods=['GET'])
@jwt_required()
@require_admin_role
def get_question_bank(qb_id):
    try:
        with db_connection() as conn:
            cur = conn.cursor()
            cur.execute("""
                SELECT qb.id, qb.original_filename, qb.description, qb.tags,
                       qb.is_active, qb.created_at, qb.parsed_questions, qb.questions_text,
                       u.name as uploaded_by_name
                FROM custom_question_bank qb
                LEFT JOIN users u ON qb.uploaded_by = u.id
                WHERE qb.id = %s
            """, (qb_id,))
            row = cur.fetchone()

        if not row:
            return jsonify({'status': 'error', 'message': 'Not found'}), 404

        parsed = _json.loads(row[6]) if isinstance(row[6], str) else (row[6] or [])

        return jsonify({'status': 'success', 'data': {
            'id': row[0], 'filename': row[1], 'description': row[2], 'tags': row[3],
            'is_active': row[4], 'created_at': str(row[5]) if row[5] else None,
            'parsed_questions': parsed, 'raw_text_preview': (row[7] or '')[:500],
            'uploaded_by': row[8], 'questions_count': len(parsed)
        }})

    except Exception as e:
        logger.error(f"[CUSTOM QB] Get failed: {e}")
        return jsonify({'status': 'error', 'message': 'Internal server error'}), 500


@admin_content_bp.route('/question-bank/<int:qb_id>', methods=['DELETE'])
@jwt_required()
@require_admin_role
def delete_question_bank(qb_id):
    try:
        with db_connection() as conn:
            cur = conn.cursor()
            cur.execute("SELECT file_path FROM custom_question_bank WHERE id = %s", (qb_id,))
            row = cur.fetchone()
            if not row:
                return jsonify({'status': 'error', 'message': 'Not found'}), 404

            filepath = row[0]
            cur.execute("DELETE FROM custom_question_bank WHERE id = %s", (qb_id,))
            conn.commit()

        if filepath and os.path.exists(filepath):
            with contextlib.suppress(Exception):
                os.remove(filepath)

        return jsonify({'status': 'success', 'message': 'Question bank deleted'})

    except Exception as e:
        logger.error(f"[CUSTOM QB] Delete failed: {e}")
        return jsonify({'status': 'error', 'message': 'Internal server error'}), 500


@admin_content_bp.route('/question-bank/<int:qb_id>/toggle', methods=['PATCH'])
@jwt_required()
@require_admin_role
def toggle_question_bank(qb_id):
    try:
        with db_connection() as conn:
            cur = conn.cursor()
            cur.execute("""
                UPDATE custom_question_bank
                SET is_active = NOT is_active, updated_at = CURRENT_TIMESTAMP
                WHERE id = %s
                RETURNING is_active
            """, (qb_id,))
            row = cur.fetchone()
            if not row:
                return jsonify({'status': 'error', 'message': 'Not found'}), 404
            conn.commit()

        return jsonify({'status': 'success', 'is_active': row[0]})

    except Exception as e:
        logger.error(f"[CUSTOM QB] Toggle failed: {e}")
        return jsonify({'status': 'error', 'message': 'Internal server error'}), 500
